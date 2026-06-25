from __future__ import annotations

import json
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "docx_visual_upgrade_work"
SOURCE = WORK / "source.docx"
OUTPUT = WORK / "HARP_VLA_Upgraded_Full_Experiment_Report_2026-06-25_VOA_visual_upgrade.docx"
DESKTOP_OUTPUT = Path("C:/Users/77941/Desktop/项目合集/HARP_VLA_Upgraded_Full_Experiment_Report_2026-06-25_VOA_visual_upgrade.docx")
FIG_DIR = ROOT / "outputs" / "voa_visual_upgrade_figures"
ROUTE_DIR = ROOT / "outputs" / "recovery_route_classifier"
TIMING_DIR = ROOT / "outputs" / "recovery_timing_ablation"


FIGURES = [
    ("fig10_decision_flow.png", "Figure V1. Representation-guided recovery decision flow. This visual summarizes how execution windows become route and timing decisions."),
    ("fig01_route_distribution.png", "Figure V2. Recovery route distribution in the current synthetic smoke run. The distribution verifies that all five routing decisions are represented."),
    ("fig02_feature_importance.png", "Figure V3. Feature importance for the recovery route classifier. Retrieval confidence, progress slope, residual, and local failure evidence are the main explanatory signals."),
    ("fig05_risk_confidence_space.png", "Figure V4. Risk-gated retrieval confidence space. Low confidence and high risk form the region where demo anchoring or human review becomes preferable."),
    ("fig06_residual_progress_space.png", "Figure V5. Residual versus progress diagnostic space. Strong recovery is associated with high action-outcome residual and low progress."),
    ("fig07_execution_manifold_proxy.png", "Figure V6. PCA proxy of the execution representation space. This is a visualization of the feature manifold, not a claim about a learned latent space unless real embeddings are supplied."),
    ("fig08_route_confidence_histogram.png", "Figure V7. Predicted route confidence distribution. This supports selective reporting and later calibration analysis."),
    ("fig09_rollout_timeline.png", "Figure V8. Example rollout timeline showing when route triggers occur during an episode."),
    ("fig03_timing_ablation.png", "Figure V9. Counterfactual timing ablation. Earlier recovery remains more effective than delayed or absent recovery under the proxy model."),
    ("fig04_timing_by_route_heatmap.png", "Figure V10. Timing sensitivity by route. Route-specific timing summaries show which cases degrade fastest when recovery is delayed."),
    ("fig11_simulation_style_diagnostics.png", "Figure V11. Simulation-style diagnostic panels. These are schematic process visuals for explanation, not real simulator screenshots."),
]


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing source docx: {SOURCE}")
    doc = Document(str(SOURCE))
    metrics = json.loads((ROUTE_DIR / "metrics.json").read_text(encoding="utf-8-sig"))
    importance = pd.read_csv(ROUTE_DIR / "feature_importance.csv")
    timing = pd.read_csv(TIMING_DIR / "summary.csv")

    add_page_break(doc)
    add_title(doc)
    add_takeaway(doc, metrics)
    add_visual_inventory(doc)
    add_pipeline_table(doc)
    add_figures(doc)
    add_metrics_table(doc, metrics)
    add_feature_table(doc, importance)
    add_timing_table(doc, timing)
    add_claims_and_limitations(doc)
    add_how_to_use(doc)

    doc.save(str(OUTPUT))
    DESKTOP_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUTPUT, DESKTOP_OUTPUT)
    print(OUTPUT)
    print(DESKTOP_OUTPUT)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_title(doc: Document) -> None:
    h = doc.add_heading("2026-06-25 Addendum: Representation-Guided Recovery Visual Upgrade", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    r = p.add_run("中文：表征引导的 VLA/VOA 操作恢复与执行期可靠性层")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(46, 94, 140)
    doc.add_paragraph(
        "This addendum integrates the 2026-06-25 VOA upgrade into the existing HARP-VLA report. "
        "The focus is visual evidence: route decisions, feature explanations, retrieval confidence, "
        "risk gating, timing ablations, and schematic execution diagnostics."
    )
    doc.add_paragraph(
        "核心表达：RAM makes the robot spatially aware before execution; HARP-VLA/VOA makes the robot failure-aware during execution."
    )
    doc.add_paragraph("中文：RAM 让机器人执行前更懂空间；HARP-VLA/VOA 让机器人执行中更懂失败。")


def add_takeaway(doc: Document, metrics: dict) -> None:
    doc.add_heading("Upgrade Summary", level=2)
    p = doc.add_paragraph()
    p.add_run("Main upgrade. ").bold = True
    p.add_run(
        "The project is no longer framed only as post-failure recovery. It now has an execution-time reliability layer that uses "
        "embedding distance, action-outcome residual, progress slope, retrieval confidence, local failure-neighbor ratio, and start risk "
        "to select the recovery route and strength."
    )
    p = doc.add_paragraph()
    p.add_run("Evidence status. ").bold = True
    p.add_run(
        "The current run is a synthetic smoke validation of the pipeline and visualization stack. It should be replaced by real VOA rollout "
        "features before making empirical robot-performance claims."
    )
    p = doc.add_paragraph()
    p.add_run("Smoke-run numbers. ").bold = True
    p.add_run(
        f"Route-classifier accuracy is {metrics['accuracy']:.3f}, macro-F1 is {metrics['macro_f1']:.3f}, "
        f"missed manual/demo escalation rate is {metrics['missed_manual_or_anchor_rate']:.3f}, and false manual/demo escalation rate is "
        f"{metrics['false_manual_or_anchor_rate']:.3f}."
    )


def add_visual_inventory(doc: Document) -> None:
    doc.add_heading("Visual Process Added", level=2)
    items = [
        "Route distribution chart: shows the five-route decision surface is populated.",
        "Feature importance chart: shows which reliability signals drive classifier decisions.",
        "Risk-confidence scatter: makes the demo-anchor and human-review gate visually inspectable.",
        "Residual-progress scatter: shows when light recovery becomes strong recovery.",
        "Execution manifold proxy: gives a compact map of success, recovery, and review regions.",
        "Route confidence histogram: supports later calibration and selective prediction.",
        "Rollout timeline: shows when recovery triggers happen during an episode.",
        "Timing ablation chart and heatmap: show immediate versus delayed recovery behavior.",
        "Simulation-style diagnostic panels: provide intuitive process visuals without claiming real simulator screenshots.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_pipeline_table(doc: Document) -> None:
    doc.add_heading("Representation-Guided Recovery Pipeline", level=2)
    rows = [
        ("Execution embedding", "policy/action/state embedding norms, embedding distance", "Places the current rollout window relative to success/recovery evidence."),
        ("Action-outcome residual", "intended action versus observed outcome", "Detects slip, contact mismatch, or failed subgoal execution."),
        ("Retrieval confidence", "retrieval distance or success/failure retrieval distances", "Separates trusted retrieval evidence from uncertain states."),
        ("Risk gate", "start risk, risk score, failure-neighbor ratio", "Routes high-risk or low-confidence cases toward demo anchor or human review."),
        ("Route classifier", "continue, recover_light, recover_strong, demo_anchor, human_review", "Turns reliability signals into recovery strength and intervention policy."),
        ("Timing ablation", "immediate, +10 steps, +40 steps, no recovery", "Tests when recovery must be triggered to remain useful."),
    ]
    add_table(doc, ["Layer", "Inputs", "Purpose"], rows)


def add_figures(doc: Document) -> None:
    doc.add_heading("Visual Evidence Pack", level=2)
    for filename, caption in FIGURES:
        path = FIG_DIR / filename
        if not path.exists():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.35))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)


def add_metrics_table(doc: Document, metrics: dict) -> None:
    doc.add_heading("Classifier Metrics", level=2)
    rows = [
        ("Run label", metrics["run_label"], "Current output label"),
        ("Dataset mode", metrics["dataset_mode"], "synthetic smoke run unless real rollout CSV is supplied"),
        ("Episodes/windows", str(metrics["n_episodes"]), "Number of rows used by the classifier"),
        ("Accuracy", f"{metrics['accuracy']:.3f}", "Overall held-out route prediction accuracy"),
        ("Macro-F1", f"{metrics['macro_f1']:.3f}", "Class-balanced route prediction score"),
        ("Missed escalation", f"{metrics['missed_manual_or_anchor_rate']:.3f}", "Manual/demo cases incorrectly routed as automatic"),
        ("False escalation", f"{metrics['false_manual_or_anchor_rate']:.3f}", "Automatic cases conservatively escalated to demo/review"),
    ]
    add_table(doc, ["Metric", "Value", "Interpretation"], rows)


def add_feature_table(doc: Document, importance: pd.DataFrame) -> None:
    doc.add_heading("Feature Importance and Interpretation", level=2)
    rows = [
        (row["feature"], f"{row['importance']:.3f}", row["interpretation"])
        for _, row in importance.sort_values("rank").iterrows()
    ]
    add_table(doc, ["Feature", "Importance", "Interpretation"], rows)


def add_timing_table(doc: Document, timing: pd.DataFrame) -> None:
    doc.add_heading("Counterfactual Recovery Timing", level=2)
    rows = [
        (
            row["policy"],
            str(row["delay_steps"]),
            f"{row['estimated_success_rate']:.3f}",
            f"{row['estimated_mean_risk']:.3f}",
            f"{row['recoverable_episode_fraction']:.3f}",
        )
        for _, row in timing.iterrows()
    ]
    add_table(doc, ["Policy", "Delay", "Est. success", "Mean risk", "Recoverable fraction"], rows)


def add_claims_and_limitations(doc: Document) -> None:
    doc.add_heading("Research Claims After This Visual Upgrade", level=2)
    sections = [
        (
            "Confirmed by the new artifact",
            [
                "The upgraded VOA/HARP-VLA reliability layer runs end to end and produces structured metrics, predictions, explanations, timing tables, and visual outputs.",
                "The method can present failure-aware execution decisions as interpretable plots rather than only text logs.",
                "The Word report now contains a richer visual appendix suitable for supervisor review, PhD application discussion, or paper-planning material.",
            ],
        ),
        (
            "Suggested but not yet proven",
            [
                "Embedding distance, residual, retrieval confidence, and risk may be sufficient to distinguish light recovery, strong recovery, demo anchoring, and human review in real rollouts.",
                "Earlier recovery should be more useful than delayed recovery when risk and residual grow over time.",
            ],
        ),
        (
            "Still unproven",
            [
                "Real simulator or hardware success-rate improvement has not been established by the synthetic smoke run.",
                "Cross-task generalization and calibrated safety thresholds still need real rollout data.",
                "The diagnostic panels are schematic explanatory visuals, not simulator screenshots.",
            ],
        ),
    ]
    for heading, items in sections:
        p = doc.add_paragraph()
        p.add_run(heading).bold = True
        for item in items:
            doc.add_paragraph(item, style="List Bullet")


def add_how_to_use(doc: Document) -> None:
    doc.add_heading("How To Replace Smoke Visuals With Real VOA Evidence", level=2)
    steps = [
        "Export a real rollout-level or window-level CSV with execution_status or route_label plus the six classifier inputs.",
        "Run: python scripts/run_voa_recovery_upgrade.py --input-rollouts path/to/voa_rollout_features.csv --run-label voa_real_rollouts",
        "Run: python scripts/generate_voa_visuals.py",
        "Rebuild this addendum so the report figures reflect the real rollout table.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "Required output locations: outputs/recovery_route_classifier/metrics.json, confusion_matrix.png, feature_importance.csv, route_predictions.csv, "
        "route_explanations.csv, and outputs/recovery_timing_ablation/summary.csv."
    )


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    doc.add_paragraph()


if __name__ == "__main__":
    main()
