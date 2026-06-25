from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "reports" / "HARP_VLA_Upgraded_Full_Experiment_Report_2026-06-25_VOA_visual_upgrade.docx"
DESKTOP_DOCX = Path("C:/Users/77941/Desktop/项目合集/HARP_VLA_Upgraded_Full_Experiment_Report_2026-06-25_VOA_visual_upgrade.docx")
FIG_DIR = ROOT / "outputs" / "voa_visual_upgrade_figures"


FIGURES = [
    ("fig10_decision_flow.png", "Front Figure 1. Representation-guided recovery decision flow."),
    ("fig07_execution_manifold_proxy.png", "Front Figure 2. Execution representation space, shown as a PCA proxy over reliability features."),
    ("fig05_risk_confidence_space.png", "Front Figure 3. Risk-gated retrieval confidence space."),
    ("fig06_residual_progress_space.png", "Front Figure 4. Action-outcome residual versus progress slope."),
    ("fig02_feature_importance.png", "Front Figure 5. Route classifier feature importance."),
    ("fig01_route_distribution.png", "Front Figure 6. Recovery route distribution."),
    ("fig08_route_confidence_histogram.png", "Front Figure 7. Predicted route confidence distribution."),
    ("fig09_rollout_timeline.png", "Front Figure 8. Example rollout trigger timeline."),
    ("fig03_timing_ablation.png", "Front Figure 9. Counterfactual recovery timing ablation."),
    ("fig04_timing_by_route_heatmap.png", "Front Figure 10. Route-level timing sensitivity heatmap."),
    ("fig11_simulation_style_diagnostics.png", "Front Figure 11. Simulation-style diagnostic panels; schematic, not real simulator screenshots."),
]


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    doc = Document(str(DOCX))
    remove_existing_front_snapshot(doc)

    insert_index = find_insert_index(doc)
    start_index = len(doc.element.body) - 1
    append_front_snapshot(doc)
    move_appended_blocks(doc, start_index, insert_index)
    remove_duplicate_snapshot_headings(doc)

    doc.save(str(DOCX))
    if DESKTOP_DOCX.parent.exists():
        shutil.copy2(DOCX, DESKTOP_DOCX)
    print(DOCX)


def remove_existing_front_snapshot(doc: Document) -> None:
    body = doc.element.body
    elements = list(body)
    start = None
    end = None
    for idx, element in enumerate(elements):
        text = "".join(element.itertext())
        if "Visual Evidence Snapshot" in text or "This front section makes the evidence visible" in text:
            start = idx
        if start is not None and "Part A - Upgraded Paper-Style Narrative" in text:
            end = idx
            break
    if start is not None and end is not None and end > start:
        for element in elements[start:end]:
            body.remove(element)


def remove_duplicate_snapshot_headings(doc: Document) -> None:
    seen = False
    for para in list(doc.paragraphs):
        if para.text.strip() != "Visual Evidence Snapshot":
            continue
        if not seen:
            seen = True
            continue
        para._element.getparent().remove(para._element)


def find_insert_index(doc: Document) -> int:
    body = doc.element.body
    for idx, element in enumerate(list(body)):
        text = "".join(element.itertext())
        if "Part A - Upgraded Paper-Style Narrative" in text:
            return idx
    for idx, element in enumerate(list(body)):
        text = "".join(element.itertext())
        if "Document Map" in text:
            return idx + 1
    return len(body) - 1


def append_front_snapshot(doc: Document) -> None:
    doc.add_heading("Visual Evidence Snapshot", level=1)
    p = doc.add_paragraph()
    p.add_run("Purpose. ").bold = True
    p.add_run(
        "This front section makes the evidence visible before the detailed narrative: "
        "runtime instability is measured first, then routed into selective recovery decisions."
    )
    p = doc.add_paragraph()
    p.add_run("Interpretation. ").bold = True
    p.add_run(
        "The figures summarize the current smoke-run pipeline: embedding drift, residual-progress diagnostics, "
        "retrieval confidence, route classification, recovery timing, and schematic execution diagnostics."
    )
    p = doc.add_paragraph()
    p.add_run("Evidence status. ").bold = True
    p.add_run(
        "These visuals validate the analysis and reporting pipeline. They should be replaced with real VOA/HARP rollout "
        "figures before making empirical robot-performance claims."
    )
    doc.add_heading("What Runtime Instability Analysis Contributes", level=2)
    p = doc.add_paragraph()
    p.add_run("Instability analysis is the evidence layer, not just a classifier. ").bold = True
    p.add_run(
        "It detects whether execution is drifting away from successful behavior, whether the commanded action is producing "
        "the expected physical outcome, whether task progress has stalled, and whether retrieved recovery evidence is still reliable."
    )
    rows = [
        ("Embedding distance", "Checks whether the current state is moving away from the success/recovery manifold."),
        ("Action-outcome residual", "Measures mismatch between intended action and observed outcome."),
        ("Progress slope", "Confirms whether the task is still advancing or has stalled."),
        ("Retrieval confidence", "Measures whether retrieved evidence is trustworthy for the current state."),
        ("Failure-neighbor ratio", "Shows whether the state is close to historical failure regions."),
        ("Route classifier", "Uses the instability evidence to choose continue, recovery strength, demo anchor, or human review."),
    ]
    add_table(doc, ["Evidence signal", "Role in instability analysis"], rows)

    add_compact_table(doc)
    for filename, caption in FIGURES:
        path = FIG_DIR / filename
        if not path.exists():
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Inches(6.15))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)
    doc.add_page_break()


def add_compact_table(doc: Document) -> None:
    rows = [
        ("Instability evidence", "embedding distance, residual, progress slope, retrieval confidence"),
        ("Decision output", "continue, recover_light, recover_strong, demo_anchor, human_review"),
        ("Timing question", "immediate recovery vs. delayed recovery vs. no recovery"),
        ("Safety interpretation", "automatic recovery when evidence is trusted; demo/human route when confidence is low"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Component"
    table.rows[0].cells[1].text = "What the teacher should see first"
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def move_appended_blocks(doc: Document, start_index: int, insert_index: int) -> None:
    body = doc.element.body
    elements = list(body)
    appended = elements[start_index:-1] if elements[-1].tag.endswith("sectPr") else elements[start_index:]
    reference = list(body)[insert_index]
    for element in appended:
        body.remove(element)
    for element in appended:
        body.insert(body.index(reference), element)


if __name__ == "__main__":
    main()
